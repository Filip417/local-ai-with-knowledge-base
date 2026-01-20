from typing import List, Optional
import os
import shutil
import re
from uuid import UUID
from app.core.database import chroma_client
from app.core.config import VECTOR_DB_COLLECTION_NAME, N_RESULTS
from app.models.message import Message
from fastapi import UploadFile
from chromadb import QueryResult
from pypdf import PdfReader
from docx import Document

data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "sources"))

def clear_documents_collection():
    """
    Deletes the collection to clear all data.
    """
    try:
        chroma_client.delete_collection(name=VECTOR_DB_COLLECTION_NAME)
        chroma_client.create_collection(name=VECTOR_DB_COLLECTION_NAME)
        return True
    except:
        return False


def add_documents(document_content: str, file_id: Optional[UUID], doc_id: Optional[str] = None):
    """
    Adds text to the vector database with file ID tracking.
    file_id is always stored in metadata for filtering; doc_id can vary per chunk.
    """
    if not document_content or not file_id:
        return False

    resolved_doc_id = doc_id or str(file_id)

    collection = chroma_client.get_or_create_collection(name=VECTOR_DB_COLLECTION_NAME)
    collection.add(
        ids=[resolved_doc_id],
        documents=[document_content],
        metadatas=[{"file_id": str(file_id)}]
    )
    return resolved_doc_id


def get_results_from_vector_db(
        last_user_message: Message,
        selected_file_ids: Optional[List[UUID]] = None) -> QueryResult | None:
    """
    Returns results from chroma db query, filtered by selected file IDs.
    """
    if not selected_file_ids:
        return None

    collection = chroma_client.get_or_create_collection(name=VECTOR_DB_COLLECTION_NAME)

    kwargs = {
        "query_texts": [last_user_message.text],
        "n_results": N_RESULTS,
    }
    file_id_strs = [str(fid) for fid in selected_file_ids]
    kwargs["where"] = {"file_id": {"$in": file_id_strs}}
    results = collection.query(**kwargs)
    return results


def save_or_reuse_data_file(file: UploadFile):
    """Save upload into backend/sources (reusing existing file name) and return the path."""
    os.makedirs(data_dir, exist_ok=True)
    raw_name = file.filename or ""
    if not raw_name:
        return False

    filename = os.path.basename(raw_name)
    destination = os.path.join(data_dir, filename)

    if os.path.exists(destination):
        return destination

    try:
        file.file.seek(0)
        with open(destination, "wb") as dest:
            shutil.copyfileobj(file.file, dest)
        return destination
    except Exception:
        return False


def upload_file_to_vector_db(file_path: str, file_extension: str, file_id: Optional[UUID] = None):
    """
    Load a saved file from disk and push its contents into the vector DB with file ID tracking.
    """
    if file_extension in ("txt", "md"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text_content = f.read()
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text_content) if p.strip()]
            for idx, paragraph in enumerate(paragraphs):
                add_documents(paragraph, file_id=file_id, doc_id=f"{file_id}-{file_extension}-{idx}")
            return True
        except Exception as e:
            print(e)
            return False
    elif file_extension == "pdf":
        try:
            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                add_documents(page.extract_text(), file_id=file_id, doc_id=f"{file_id}-pdf-{idx}")
            return True
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return False
    elif file_extension == "docx":
        try:
            doc = Document(file_path)
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    add_documents(para.text, file_id=file_id, doc_id=f"{file_id}-docx-p-{idx}")
            
            # Also extract text from tables
            for t_idx, table in enumerate(doc.tables):
                table_content = ""
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_content += cell.text + "\n\n"
                if table_content.strip():
                    add_documents(table_content, file_id=file_id, doc_id=f"{file_id}-docx-t-{t_idx}")
            
            return True
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return False
    return False


def delete_file_from_disk(filename: str):
    """Delete a file from the backend/sources directory."""
    try:
        file_path = os.path.join(data_dir, os.path.basename(filename))
        
        if os.path.exists(file_path):
            os.remove(file_path)
        return True
    except Exception as e:
        print(f"Error deleting file from disk: {e}")
        return False


def delete_file_from_vector_db(file_id: UUID):
    """
    Delete documents associated with a file ID from the vector database.
    """
    try:
        collection = chroma_client.get_or_create_collection(name=VECTOR_DB_COLLECTION_NAME)
        
        # Delete documents by file ID
        collection.delete(ids=[str(file_id)])
        
        return True
    except Exception as e:
        print(f"Error deleting file from vector db: {e}")
        return False


def handle_file_stream(file):
    if file.extension in ('txt', 'md'):
        with open(file.path, 'r', encoding='utf-8') as f:
            yield from f
    elif file.extension == "pdf":
        try:
            reader = PdfReader(file.path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    yield text + "\n\n"
        except Exception as e:
            yield f"Error reading PDF: {str(e)}"
    elif file.extension == "docx":
        try:
            doc = Document(file.path)
            for para in doc.paragraphs:
                if para.text.strip():
                    yield para.text + "\n\n"
            
            # Also yield text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            yield cell.text + "\n\n"
        except Exception as e:
            yield f"Error reading DOCX: {str(e)}"